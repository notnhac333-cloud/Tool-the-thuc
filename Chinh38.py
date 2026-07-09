import sys
import os
import re
import copy
import time
import subprocess
import datetime
import pythoncom
import requests
import json

from PyQt6.QtWidgets import (QApplication, QMainWindow, QLabel, QVBoxLayout, 
                             QWidget, QMessageBox, QHBoxLayout, QPushButton, 
                             QFileDialog, QScrollArea, QProgressBar, QSplitter, QCheckBox)
from PyQt6.QtCore import Qt, QUrl, QThread, pyqtSignal
from PyQt6.QtGui import QDragEnterEvent, QDropEvent, QDesktopServices, QImage, QPixmap

from docx import Document
from docx.shared import Mm, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# --- TRÌNH XEM TRƯỚC (PREVIEW) TÍCH HỢP ZOOM BẰNG CHUỘT ---
class PDFViewer(QWidget):
    def __init__(self):
        super().__init__()
        self.zoom_factor = 1.0
        self.original_pixmaps = []
        self.labels = []
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # Thanh công cụ Zoom
        toolbar = QHBoxLayout()
        self.btn_zoom_out = QPushButton("🔍 Thu nhỏ (-)")
        self.btn_zoom_in = QPushButton("🔍 Phóng to (+)")
        self.btn_fit = QPushButton("🖵 Vừa màn hình")
        
        for btn in [self.btn_zoom_out, self.btn_zoom_in, self.btn_fit]:
            btn.setStyleSheet("background: #ecf0f1; border: 1px solid #bdc3c7; border-radius: 4px; padding: 5px 10px; font-weight: bold; color: #2c3e50;")
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            
        toolbar.addWidget(self.btn_zoom_out)
        toolbar.addWidget(self.btn_zoom_in)
        toolbar.addWidget(self.btn_fit)
        toolbar.addStretch()
        layout.addLayout(toolbar)
        
        # Khu vực cuộn hiển thị ảnh
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setStyleSheet("background-color: #bdc3c7; border: none;")
        
        self.container = QWidget()
        self.container.setStyleSheet("background-color: #bdc3c7;")
        self.container_layout = QVBoxLayout(self.container)
        self.container_layout.setAlignment(Qt.AlignmentFlag.AlignHCenter | Qt.AlignmentFlag.AlignTop)
        
        self.scroll_area.setWidget(self.container)
        layout.addWidget(self.scroll_area)
        
        self.btn_zoom_in.clicked.connect(self.zoom_in)
        self.btn_zoom_out.clicked.connect(self.zoom_out)
        self.btn_fit.clicked.connect(self.fit_width)
        
        # Lắng nghe sự kiện Ctrl + Cuộn chuột
        self.scroll_area.viewport().installEventFilter(self)

    def eventFilter(self, source, event):
        if event.type() == event.Type.Wheel and source is self.scroll_area.viewport():
            modifiers = QApplication.keyboardModifiers()
            if modifiers == Qt.KeyboardModifier.ControlModifier:
                if event.angleDelta().y() > 0:
                    self.zoom_in()
                else:
                    self.zoom_out()
                return True
        return super().eventFilter(source, event)
        
    def show_message(self, msg):
        self.clear_view()
        lbl = QLabel(msg)
        lbl.setStyleSheet("color: #2c3e50; font-size: 14px; margin-top: 20px;")
        lbl.setAlignment(Qt.AlignmentFlag.AlignHCenter)
        self.container_layout.addWidget(lbl)
        
    def clear_view(self):
        self.original_pixmaps.clear()
        self.labels.clear()
        for i in reversed(range(self.container_layout.count())): 
            widget = self.container_layout.itemAt(i).widget()
            if widget: widget.setParent(None)

    def load_pdf(self, pdf_path):
        self.clear_view()
        if not pdf_path or not os.path.exists(pdf_path):
            self.show_message("Lỗi: Không tìm thấy file hiển thị xem trước.")
            return

        try:
            import fitz
            doc = fitz.open(pdf_path)
            for i in range(len(doc)):
                page = doc.load_page(i)
                pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
                fmt = QImage.Format.Format_RGBA8888 if pix.alpha else QImage.Format.Format_RGB888
                img = QImage(pix.samples, pix.width, pix.height, pix.stride, fmt)
                pixmap = QPixmap.fromImage(img)
                self.original_pixmaps.append(pixmap)
                
                lbl = QLabel()
                lbl.setStyleSheet("background-color: white; margin-bottom: 15px;")
                self.labels.append(lbl)
                self.container_layout.addWidget(lbl)
            doc.close()
            self.fit_width()
            try: os.remove(pdf_path)
            except: pass
        except ImportError:
            self.show_message("<b>Chưa cài thư viện render đồ họa!</b><br>Mở PowerShell gõ lệnh:<br>pip install PyMuPDF")
        except Exception as e:
            self.show_message(f"Lỗi hiển thị PDF: {str(e)}")

    def apply_zoom(self):
        if not self.original_pixmaps: return
        for i, pixmap in enumerate(self.original_pixmaps):
            new_w = int(pixmap.width() * self.zoom_factor)
            scaled = pixmap.scaledToWidth(new_w, Qt.TransformationMode.SmoothTransformation)
            self.labels[i].setPixmap(scaled)

    def zoom_in(self):
        self.zoom_factor += 0.15
        self.apply_zoom()

    def zoom_out(self):
        self.zoom_factor -= 0.15
        if self.zoom_factor < 0.2: self.zoom_factor = 0.2
        self.apply_zoom()

    def fit_width(self):
        if not self.original_pixmaps: return
        vp_width = self.scroll_area.viewport().width() - 40
        if vp_width <= 0: vp_width = 800
        base_w = self.original_pixmaps[0].width()
        self.zoom_factor = vp_width / base_w
        self.apply_zoom()


class TuDienLocal:
    @staticmethod
    def load_dict(log_callback=None):
        tu_dien = set()
        base_dir1 = os.path.dirname(os.path.abspath(__file__))
        base_dir2 = os.getcwd()
        try: base_dir3 = os.path.dirname(os.path.realpath(sys.argv[0]))
        except: base_dir3 = base_dir1
            
        possible_folders = [
            os.path.join(base_dir1, 'dict'), os.path.join(base_dir1, 'dist'), base_dir1,
            os.path.join(base_dir2, 'dict'), os.path.join(base_dir2, 'dist'), base_dir2,
            os.path.join(base_dir3, 'dict'), os.path.join(base_dir3, 'dist'), base_dir3
        ]

        for folder in possible_folders:
            if os.path.exists(folder) and os.path.isdir(folder):
                for file_name in os.listdir(folder):
                    if file_name.lower().endswith(('.dic', '.txt', '.dict')):
                        if "api_key" in file_name.lower(): continue
                        file_path = os.path.join(folder, file_name)
                        for enc in ['utf-8', 'utf-16', 'cp1252']:
                            try:
                                with open(file_path, "r", encoding=enc) as f:
                                    for line in f:
                                        w = line.strip().lower().split('/')[0]
                                        if w and not w.isnumeric() and not w.startswith("["):
                                            tu_dien.add(w)
                                break
                            except Exception: pass
                if len(tu_dien) > 0: break
                
        if len(tu_dien) > 0:
            if log_callback: log_callback(f"<i>➜ Đã nạp Từ điển Local ({len(tu_dien)} âm tiết).</i>", "#27ae60")
        return tu_dien

class KiemTraChinhTaLocalAI:
    @staticmethod
    def quet_loi(text_full, log_callback, progress_callback, p_start, p_end):
        paragraphs = [p.strip() for p in text_full.split('\n') if len(p.strip()) > 5]
        chunks = []
        current_chunk = ""
        for p in paragraphs:
            if len(current_chunk) + len(p) < 1200:
                current_chunk += p + "\n"
            else:
                chunks.append(current_chunk.strip())
                current_chunk = p + "\n"
        if current_chunk.strip(): chunks.append(current_chunk.strip())

        loi_phat_hien = []
        danh_sach_tu_sai = []
        total_chunks = len(chunks)
        chunk_step = (p_end - p_start) / max(1, total_chunks)
        current_p = p_start

        url = "http://localhost:11434/api/generate"
        
        # THUẬT TOÁN FALLBACK THÔNG MINH CHO CẤU HÌNH MÁY
        models_to_try = ["qwen3:14b", "qwen2.5:14b", "qwen2.5:7b"]
        installed_models = []
        try:
            res = requests.get("http://localhost:11434/api/tags", timeout=3)
            if res.status_code == 200:
                installed_models = [m["name"] for m in res.json().get("models", [])]
        except: pass

        available_models = [m for m in models_to_try if any(m in im for im in installed_models)]
        if not available_models: available_models = ["qwen3:14b"] 

        for idx, chunk in enumerate(chunks):
            log_callback(f"<i>➜ AI đang soi kỹ đoạn {idx+1}/{total_chunks}...</i>", "#7f8c8d")
            prompt = f"""Hãy đóng vai một chuyên gia biên tập văn bản khắt khe nhất.
Nhiệm vụ: Đọc đoạn văn sau và tìm MỌI lỗi gõ sai chữ, dư chữ, sai chính tả hoặc sai ngữ cảnh (ví dụ: 'càn bộ' -> 'cán bộ', 'báo cá' -> 'báo cáo', 'Nnữ' -> 'Nữ', 'đồn' -> 'đồng').
Chỉ trả về các lỗi theo đúng mẫu sau (mỗi lỗi 1 dòng):
LỖI: [từ gõ sai] -> SỬA THÀNH: [từ đúng]
Tuyệt đối không giải thích. Nếu không có lỗi, in ra duy nhất chữ: OK

Đoạn văn:
'''{chunk}'''"""
            
            success = False
            for target_model in available_models:
                payload = {"model": target_model, "prompt": prompt, "stream": False, "options": {"temperature": 0.0}}
                try:
                    response = requests.post(url, json=payload, timeout=90)
                    response.raise_for_status()
                    res_text = response.json().get("response", "").strip()
                    success = True
                    if idx == 0: log_callback(f"<i>➜ Đã khóa mục tiêu vào model tự động nhận diện: <b>{target_model}</b></i>", "#8e44ad")
                    
                    for line in res_text.split('\n'):
                        line_up = line.upper()
                        if "LỖI:" in line_up and "SỬA THÀNH:" in line_up:
                            try:
                                match = re.search(r'(?i)LỖI:\s*(.*?)\s*->\s*(?i)SỬA THÀNH:\s*(.*)', line)
                                if match:
                                    sai = match.group(1).strip(' \'"“”.,;:')
                                    dung = match.group(2).strip(' \'"“”.,;:')
                                    if len(sai) >= 2 and sai.lower() not in [t.lower() for t in danh_sach_tu_sai]:
                                        danh_sach_tu_sai.append(sai)
                                        loi_phat_hien.append(f"- Lỗi ngữ cảnh (AI): <b><span style='color:red;'>{sai}</span></b> ➜ Sửa thành: <b><span style='color:#27ae60;'>{dung}</span></b>")
                            except: pass
                    break 
                except requests.exceptions.ConnectionError:
                    log_callback(f"<i>➜ Cảnh báo: Ollama chưa bật. Hủy bỏ phân tích AI!</i>", "#e67e22")
                    progress_callback(p_end)
                    return [], []
                except Exception as e:
                    log_callback(f"<i>➜ Model {target_model} bị lỗi/quá tải. Đang tự động lùi xuống model nhẹ hơn...</i>", "#d35400")
                    continue
            
            if not success:
                log_callback(f"<i>➜ Bỏ qua phân tích đoạn {idx+1} do không có phản hồi AI.</i>", "#e74c3c")

            current_p += chunk_step
            progress_callback(current_p)

        danh_sach_tu_sai.sort(key=len, reverse=True)
        return loi_phat_hien, danh_sach_tu_sai

class TienIchDocx:
    TU_KHOA_TIEU_DE = ["BÁO CÁO", "TỜ TRÌNH", "KẾ HOẠCH", "QUYẾT ĐỊNH", "ĐỀ CƯƠNG", "CHUYÊN ĐỀ", "THÔNG BÁO", "CHỈ THỊ", "QUY CHẾ", "QUY ĐỊNH", "HƯỚNG DẪN", "CHƯƠNG TRÌNH", "PHƯƠNG ÁN", "KẾT LUẬN", "BIÊN BẢN", "NGHỊ QUYẾT"]

    @staticmethod
    def xoa_element(element):
        if element is not None:
            parent = element.getparent()
            if parent is not None: parent.remove(element)

    @staticmethod
    def thiet_lap_font(run, size=14, bold=False, italic=False, underline=False, color_rgb=None):
        run.font.name = 'Times New Roman'
        if size is not None: run.font.size = Pt(size)
        if bold is not None: run.font.bold = bold
        if italic is not None: run.font.italic = italic
        if underline: run.underline = True
        if color_rgb: run.font.color.rgb = color_rgb
        if run._element.rPr is not None and run._element.rPr.rFonts is not None:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    @staticmethod
    def sao_chep_run(paragraph, run_old, text_to_insert, color_rgb=None, force_size=None, force_bold=None, force_italic=None, force_underline=None):
        r_new = paragraph.add_run(text_to_insert)
        r_new.font.name = 'Times New Roman'
        if force_size is not None: r_new.font.size = Pt(force_size)
        elif run_old and run_old.font.size: r_new.font.size = run_old.font.size
        else: r_new.font.size = Pt(14)
        r_new.font.bold = force_bold if force_bold is not None else (run_old.font.bold if run_old else False)
        r_new.font.italic = force_italic if force_italic is not None else (run_old.font.italic if run_old else False)
        r_new.font.underline = force_underline if force_underline is not None else (run_old.font.underline if run_old else False)
        if run_old and run_old.font.subscript: r_new.font.subscript = True
        if run_old and run_old.font.superscript: r_new.font.superscript = True
        if color_rgb: r_new.font.color.rgb = color_rgb
        if r_new._element.rPr is not None and r_new._element.rPr.rFonts is not None:
            r_new._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        return r_new

    @staticmethod
    def highlight_paragraph(p, danh_sach_tu_sai):
        if not danh_sach_tu_sai: return
        danh_sach_hop_le = [tu.strip() for tu in danh_sach_tu_sai if tu.strip()]
        if not danh_sach_hop_le: return

        merged_runs_info = []
        for r in p.runs:
            if not r.text: continue
            fmt = {
                'bold': r.font.bold, 'italic': r.font.italic, 'underline': r.font.underline,
                'size': r.font.size, 'name': r.font.name or 'Times New Roman',
                'color': r.font.color.rgb if r.font.color else None,
                'subscript': r.font.subscript, 'superscript': r.font.superscript
            }
            if not merged_runs_info:
                merged_runs_info.append({'text': r.text, 'fmt': fmt})
            else:
                prev_fmt = merged_runs_info[-1]['fmt']
                if (fmt['bold'] == prev_fmt['bold'] and fmt['italic'] == prev_fmt['italic'] and 
                    fmt['underline'] == prev_fmt['underline'] and fmt['subscript'] == prev_fmt['subscript'] and 
                    fmt['superscript'] == prev_fmt['superscript'] and fmt['color'] == prev_fmt['color']):
                    merged_runs_info[-1]['text'] += r.text
                else:
                    merged_runs_info.append({'text': r.text, 'fmt': fmt})
        
        for r in p.runs: r.clear()
        
        escaped_words = [re.escape(w) for w in danh_sach_hop_le]
        vi_chars = r'a-zA-ZđĐáàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵ'
        pattern = rf'(?i)(?<![{vi_chars}])({"|".join(escaped_words)})(?![{vi_chars}])'
        
        for info in merged_runs_info:
            text = info['text']
            fmt = info['fmt']
            
            parts = re.split(pattern, text)
            for part in parts:
                if not part: continue
                r_new = p.add_run(part)
                r_new.font.name = fmt['name']
                if fmt['size']: r_new.font.size = fmt['size']
                r_new.font.italic = fmt['italic']
                r_new.font.underline = fmt['underline']
                r_new.font.subscript = fmt['subscript']
                r_new.font.superscript = fmt['superscript']
                
                is_wrong = False
                for w in danh_sach_hop_le:
                    if part.lower() == w.lower():
                        is_wrong = True
                        break
                
                if is_wrong:
                    r_new.font.bold = True
                    r_new.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    r_new.font.bold = fmt['bold']
                    if fmt['color']: r_new.font.color.rgb = fmt['color']

    @staticmethod
    def copy_text_only_with_subscript(paragraph, text, size=14, bold=False, italic=False, underline=False, color_rgb=None):
        parts = re.split(r'([A-Za-zĐđ]+)(\d+)', text)
        i = 0
        while i < len(parts):
            if not parts[i]:
                i += 1; continue
            if i + 2 < len(parts) and re.match(r'^[A-Za-zĐđ]+$', parts[i+1]) and re.match(r'^\d+$', parts[i+2]):
                r_new = paragraph.add_run(parts[i] + parts[i+1])
                TienIchDocx.thiet_lap_font(r_new, size, bold, italic, underline, color_rgb)
                r_sub = paragraph.add_run(parts[i+2])
                TienIchDocx.thiet_lap_font(r_sub, size, bold, italic, underline, color_rgb)
                r_sub.font.subscript = True
                i += 3
            else:
                r_new = paragraph.add_run(parts[i])
                TienIchDocx.thiet_lap_font(r_new, size, bold, italic, underline, color_rgb)
                i += 1

    @staticmethod
    def ve_duong_ke_shape(paragraph, do_dai_pt):
        xml = f"""
        <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:v="urn:schemas-microsoft-com:vml">
            <w:pict>
                <v:line style="position:absolute; z-index:1; mso-position-horizontal:center; mso-position-horizontal-relative:margin; margin-top:17pt;"
                        from="0,0" to="{do_dai_pt}pt,0" strokecolor="#000000" strokeweight="0.75pt" />
            </w:pict>
        </w:r>
        """
        paragraph._p.append(parse_xml(xml))

    @staticmethod
    def tao_dong_trang_chu_ky(cell_obj, so_dong, size=14):
        for _ in range(so_dong):
            p = cell_obj.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(0), Pt(0)
            p.paragraph_format.widow_control = False
            run = p.add_run("\u00A0")
            TienIchDocx.thiet_lap_font(run, size)

    @staticmethod
    def add_blank_line_if_needed(doc_or_cell, size=14):
        if hasattr(doc_or_cell, 'paragraphs') and len(doc_or_cell.paragraphs) > 0:
            last_p = doc_or_cell.paragraphs[-1]
            if not last_p.text.strip() and '<w:pict>' not in last_p._element.xml:
                return 
        p = doc_or_cell.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(" ")
        TienIchDocx.thiet_lap_font(r, size)

    @staticmethod
    def convert_doc_to_docx(doc_path):
        try:
            import win32com.client as win32
            docx_path = doc_path + "x"
            if os.path.exists(docx_path): os.remove(docx_path)
            word = win32.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            doc = word.Documents.Open(os.path.abspath(doc_path), False, True)
            doc.SaveAs2(os.path.abspath(docx_path), FileFormat=16)
            doc.Close(SaveChanges=False)
            word.Quit()
            return docx_path
        except Exception as e:
            try: word.Quit()
            except: pass
            raise Exception(f"Lỗi convert .doc -> .docx: {str(e)}")

class BoXuLyVanBan:
    @staticmethod
    def phan_loai_van_ban(doc_cu):
        for p in doc_cu.paragraphs[:20]:
            if "ĐẢNG CỘNG SẢN" in p.text.upper(): return "HD36"
        for table in doc_cu.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "ĐẢNG CỘNG SẢN" in cell.text.upper(): return "HD36"
        return "ND30"

    @staticmethod
    def is_invisible_table(tbl):
        tblPr = tbl._tbl.tblPr
        borders = tblPr.xpath('./w:tblBorders')
        if not borders: return True
        border_nodes = borders[0].xpath('./*')
        if not border_nodes: return True
        for edge in border_nodes:
            val = edge.get(qn('w:val'))
            if val and val.lower() not in ['none', 'nil', '0']:
                return False
        return True

    @staticmethod
    def get_text_convert_bullet(p_old):
        text = p_old.text.strip()
        numPrs = p_old._element.xpath('.//*[local-name()="numPr"]')
        if numPrs and text:
            if not text.startswith('-') and not text.startswith('+'):
                text = "- " + text
        return text

    @staticmethod
    def bat_loi_hanh_chinh(text_full):
        loi = []
        tu_sai = []
        matches = re.finditer(r'(?i)(\d[\d\.\,]*)(?:\s|\xa0)*(đồn|đồg|đg|đng)(?![a-zA-ZđĐáàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵ])', text_full)
        for m in matches:
            sai = m.group(0)
            dung = m.group(1) + " đồng"
            loi.append(f"- Lỗi đánh máy tiền tệ: <b><span style='color:red;'>{sai}</span></b> ➜ Sửa chuẩn: <b><span style='color:#27ae60;'>{dung}</span></b>")
            tu_sai.append(m.group(2))
        return loi, tu_sai

    @staticmethod
    def export_preview_pdf(docx_path):
        pdf_path = docx_path.replace(".docx", "_preview.pdf").replace(".doc", "_preview.pdf")
        word = None
        doc = None
        try:
            pythoncom.CoInitialize()
            import win32com.client as win32
            word = win32.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            doc = word.Documents.Open(os.path.abspath(docx_path), False, True)
            doc.SaveAs2(os.path.abspath(pdf_path), FileFormat=17) 
            doc.Close(SaveChanges=False)
            word.Quit()
            word = None
            return pdf_path
        except Exception:
            if doc: 
                try: doc.Close(SaveChanges=False)
                except: pass
            if word: 
                try: word.Quit()
                except: pass
            return None
        finally:
            pythoncom.CoUninitialize()

    @staticmethod
    def process_single_pass(input_path, output_path, log_callback, progress_callback, enable_spellcheck=True, shrink_mode=False):
        doc_cu = Document(input_path)
        loai_vb = BoXuLyVanBan.phan_loai_van_ban(doc_cu)
        
        log_callback(f"<i>➔ Cấu trúc bám sát: <b>{'Hướng dẫn 36 (Đảng)' if loai_vb == 'HD36' else 'Nghị định 30 (Chính quyền)'}</b></i>", "#3498db")
        progress_callback(5.0)

        danh_sach_tu_sai_tong = []

        if not shrink_mode:
            full_text_to_check = []
            for p in doc_cu.paragraphs:
                if p.text.strip(): full_text_to_check.append(p.text.strip())
            for tbl in doc_cu.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        if cell.text.strip(): full_text_to_check.append(cell.text.strip())
            
            text_for_ai = "\n".join(full_text_to_check)
            if text_for_ai:
                progress_callback(10.0)
                
                if enable_spellcheck:
                    loi_rg, tu_rg = BoXuLyVanBan.bat_loi_hanh_chinh(text_for_ai)
                    if loi_rg:
                        danh_sach_tu_sai_tong.extend(tu_rg)
                        for l in loi_rg: log_callback(l, "#e74c3c")
                    
                    loi_ol, tu_ol = KiemTraChinhTaLocalAI.quet_loi(text_for_ai, log_callback, progress_callback, 10.0, 60.0)
                    if loi_ol:
                        danh_sach_tu_sai_tong.extend(tu_ol)
                        log_callback(f"<b>AI phát hiện {len(loi_ol)} lỗi đánh máy/ngữ cảnh:</b>", "#c0392b")
                        for l in loi_ol: log_callback(l, "#e74c3c")

                    if not loi_rg and not loi_ol:
                        log_callback("<i>➜ Hoàn hảo: Bộ AI không phát hiện từ sai ngữ cảnh/chính tả nào!</i>", "#27ae60")
                else:
                    log_callback("<i>➜ BỎ QUA SOÁT LỖI. Chỉ tự động format chuẩn Nghị định 30 / HD 36...</i>", "#2980b9")
        
        progress_callback(65.0)
        log_callback("<i>➔ Đang tạo bản nháp mới và tái tạo 100% định dạng...</i>", "#7f8c8d")

        if loai_vb == "ND30":
            title_size, kg_size, body_size, sig_size, name_size = (13 if shrink_mode else 14,) * 4 + (14,)
        else:
            title_size, kg_size, body_size, sig_size, name_size = 15, (13 if shrink_mode else 14), (13 if shrink_mode else 14), 14, 14

        doc_moi = Document()
        style = doc_moi.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(body_size)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.widow_control = False 
        
        if style._element.rPr.rFonts is None:
            rFonts = OxmlElement('w:rFonts'); style._element.rPr.append(rFonts)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        if len(doc_moi.paragraphs) > 0:
            TienIchDocx.xoa_element(doc_moi.paragraphs[0]._element)

        section = doc_moi.sections[0]
        section.page_height, section.page_width = Mm(297), Mm(210)
        section.top_margin, section.bottom_margin = Mm(20), Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(15) if loai_vb == "HD36" else Mm(20)
        section.different_first_page_header_footer = True
        section.header_distance = Cm(1.0) if loai_vb == "HD36" else Cm(1.27)
        section.footer_distance = Cm(1.0) if loai_vb == "HD36" else Cm(1.27)
        
        header = section.header
        p_head = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_head.text = ""
        run_head = p_head.add_run()
        TienIchDocx.thiet_lap_font(run_head, 14)
        for el_type, val in [('begin', None), (None, 'PAGE   \\* MERGEFORMAT'), ('separate', None), (None, '2'), ('end', None)]:
            if el_type:
                fld = OxmlElement('w:fldChar')
                fld.set(qn('w:fldCharType'), el_type)
                run_head._r.append(fld)
            else:
                if 'PAGE' in val:
                    instr = OxmlElement('w:instrText')
                    instr.set(qn('xml:space'), 'preserve')
                    instr.text = val
                    run_head._r.append(instr)
                else:
                    t = OxmlElement('w:t'); t.text = val
                    run_head._r.append(t)
        for p in section.first_page_header.paragraphs: p.text = ""

        header_table = None
        footer_table = None
        
        if doc_cu.tables:
            for tbl in doc_cu.tables:
                tbl_text = "".join(c.text for row in tbl.rows for c in row.cells).upper()
                if not header_table and ("CỘNG HÒA" in tbl_text or "ĐẢNG CỘNG SẢN" in tbl_text or "ỦY BAN" in tbl_text):
                    header_table = tbl
            for tbl in reversed(doc_cu.tables):
                tbl_text = "".join(c.text for row in tbl.rows for c in row.cells).upper()
                if "NƠI NHẬN" in tbl_text and ("CHỦ TỊCH" in tbl_text or "TM." in tbl_text or "BÍ THƯ" in tbl_text):
                    footer_table = tbl
                    break

        progress_callback(70.0)

        pre_footer_elements = []
        post_footer_elements = []
        passed_footer = False

        for element in doc_cu.element.body:
            if element.tag.endswith('tbl'):
                from docx.table import Table
                tbl = Table(element, doc_cu)
                if footer_table and tbl._element == footer_table._element:
                    passed_footer = True
                    continue
                if header_table and tbl._element == header_table._element: continue
                if passed_footer:
                    post_footer_elements.append(('tbl', tbl))
                    continue

                tbl_text = "".join(c.text for row in tbl.rows for c in row.cells).upper()
                if "KÍNH GỬI" in tbl_text and len(tbl_text) < 600 and BoXuLyVanBan.is_invisible_table(tbl):
                    for row in tbl.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.text.strip(): pre_footer_elements.append(('p', p))
                else:
                    pre_footer_elements.append(('tbl', tbl))
            
            elif element.tag.endswith('p'):
                from docx.text.paragraph import Paragraph
                p = Paragraph(element, doc_cu)
                if p.text.strip() or len(p.runs) > 0:
                    if passed_footer: post_footer_elements.append(('p', p))
                    else: pre_footer_elements.append(('p', p))

        is_cong_van = True
        for item_type, item in pre_footer_elements[:5]:
            if item_type == 'p':
                tu = item.text.strip().upper()
                if any(tu.startswith(kw) for kw in TienIchDocx.TU_KHOA_TIEU_DE):
                    is_cong_van = False; break

        kinh_gui_buffer = []
        i = 0
        while i < len(pre_footer_elements):
            item_type, item = pre_footer_elements[i]
            if item_type == 'p':
                text = BoXuLyVanBan.get_text_convert_bullet(item).strip()
                tu = text.upper()
                if tu.startswith("KÍNH GỬI"):
                    kinh_gui_buffer.append(item)
                    pre_footer_elements.pop(i)
                    prev_text = text
                    
                    if text.endswith(".") and not text.endswith("..."): break
                    while i < len(pre_footer_elements):
                        next_type, next_item = pre_footer_elements[i]
                        if next_type == 'tbl': break 
                        next_text = BoXuLyVanBan.get_text_convert_bullet(next_item).strip()
                        if not next_text:
                            pre_footer_elements.pop(i)
                            continue
                        next_tu = next_text.upper()
                        if any(next_tu.startswith(kw) for kw in ["CĂN CỨ", "THỰC HIỆN", "THEO ", "NHẰM ", "BÁO CÁO", "A.", "I.", "SAU KHI", "QUYẾT ĐỊNH", "ĐIỀU 1", "ĐIỀU I"]) or len(next_text) > 150:
                            break
                        if not (next_text.startswith("-") or next_text.startswith("+") or next_tu.startswith("KÍNH GỬI")):
                            if not (prev_text.endswith(";") or prev_text.endswith(":") or prev_text.endswith(",")):
                                break
                        kinh_gui_buffer.append(next_item)
                        pre_footer_elements.pop(i)
                        prev_text = next_text
                        if next_text.endswith("."): break
                    break
            i += 1

        body_start_idx = 0
        if not is_cong_van:
            body_start_idx = len(pre_footer_elements)
            title_found = False
            for i, (item_type, item) in enumerate(pre_footer_elements):
                if item_type == 'tbl':
                    body_start_idx = i; break
                text = item.text.strip()
                tu = text.upper()
                if not text: continue
                if not title_found:
                    title_found = True; continue 
                if tu.startswith("V/V") or tu.startswith("VỀ VIỆC"): continue 
                if any(tu.startswith(kw) for kw in ["CĂN CỨ", "THỰC HIỆN", "THEO ", "NHẰM ", "XÉT ", "QUA ", "HIỆN NAY", "SAU KHI", "ĐIỀU 1", "ĐIỀU I"]) or re.match(r'^([A-ZIVX]+\.|[0-9]+\.)\s+[A-ZĐa-zđ]', text) or len(text) > 130:
                    body_start_idx = i; break

        progress_callback(75.0)

        if header_table:
            valid_rows = [row for row in header_table.rows if "".join(cell.text for cell in row.cells).strip()]
            left_lines_data = [] 
            right_lines = []
            for row in valid_rows:
                if len(row.cells) >= 2:
                    for p in row.cells[0].paragraphs:
                        text = p.text.strip()
                        if text: left_lines_data.append(text)
                    right_lines.extend([l.strip() for l in row.cells[-1].text.split('\n') if l.strip()])

            co_quan_lines = []; so_lines = []
            for text in left_lines_data:
                lu = text.upper()
                if lu.startswith("SỐ") or "SỐ:" in lu or lu.startswith("V/V") or "VỀ VIỆC" in lu: so_lines.append(text)
                else: co_quan_lines.append(text)

            quoc_hieu_lines = []; ngay_lines = []
            for l in right_lines:
                lu = l.upper()
                if "NGÀY" in lu and "THÁNG" in lu and "NĂM" in lu: ngay_lines.append(l)
                else: quoc_hieu_lines.append(l)

            new_table = doc_moi.add_table(rows=2, cols=2)
            new_table.autofit = False
            tblPr = new_table._tbl.tblPr
            for w in tblPr.xpath('./w:tblW'): tblPr.remove(w)
            for layout in tblPr.xpath('./w:tblLayout'): tblPr.remove(layout)
            
            layout_fixed = OxmlElement('w:tblLayout')
            layout_fixed.set(qn('w:type'), 'fixed')
            tblPr.append(layout_fixed)

            if loai_vb == "ND30":
                new_table.columns[0].width = Mm(60)
                new_table.columns[1].width = Mm(100)
                for row in new_table.rows:
                    row.cells[0].width = Mm(60); row.cells[1].width = Mm(100)
            else:
                new_table.columns[0].width = Mm(80)
                new_table.columns[1].width = Mm(85)

            for row in new_table.rows:
                for c in row.cells:
                    for p in list(c.paragraphs): TienIchDocx.xoa_element(p._element)

            cell_tl = new_table.cell(0, 0)
            for i, line in enumerate(co_quan_lines):
                p = cell_tl.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.widow_control = False
                p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(0), Pt(0)
                
                lu_check = line.upper().replace("UỶ", "ỦY")
                if "*" in lu_check and loai_vb == "HD36":
                    TienIchDocx.sao_chep_run(p, None, "*", force_size=14, force_bold=False)
                    continue
                
                is_bold = True if ("ỦY BAN NHÂN DÂN" in lu_check or "XÃ ĐỒNG TIẾN" in lu_check or i == len(co_quan_lines)-1) else False
                TienIchDocx.sao_chep_run(p, None, lu_check, force_size=13 if loai_vb == "ND30" else 14, force_bold=is_bold)
                
                if i == len(co_quan_lines) - 1 and loai_vb == "ND30":
                    TienIchDocx.ve_duong_ke_shape(p, 48)

            if loai_vb == "ND30":
                TienIchDocx.add_blank_line_if_needed(new_table.cell(1, 0), size=13)

            cell_bl = new_table.cell(1, 0)
            is_ty_cv = False
            for line in so_lines:
                p = cell_bl.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.widow_control = False
                p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(0), Pt(0)
                line = line.strip('\n\r') 
                
                if line.upper().startswith("V/V") or "VỀ VIỆC" in line.upper():
                    is_ty_cv = True
                    TienIchDocx.copy_text_only_with_subscript(p, line, 13, bold=False, italic=False)
                elif is_ty_cv:
                    TienIchDocx.copy_text_only_with_subscript(p, line, 13, bold=False, italic=False)
                else:
                    TienIchDocx.copy_text_only_with_subscript(p, line, 13 if loai_vb=="ND30" else 14, bold=False)

            cell_tr = new_table.cell(0, 1)
            for i, line in enumerate(quoc_hieu_lines):
                p = cell_tr.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.widow_control = False
                p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(0), Pt(0)
                lu = line.upper()
                if "ĐẢNG CỘNG SẢN" in lu:
                    TienIchDocx.sao_chep_run(p, None, lu, force_size=15, force_bold=True)
                    if loai_vb == "HD36": TienIchDocx.ve_duong_ke_shape(p, 206)
                elif "CỘNG HÒA" in lu or "CỘNG HOÀ" in lu:
                    TienIchDocx.sao_chep_run(p, None, lu, force_size=13, force_bold=True)
                elif "ĐỘC LẬP" in lu:
                    line = re.sub(r'\s*[-–—]\s*', ' - ', line)
                    TienIchDocx.sao_chep_run(p, None, line, force_size=14, force_bold=True)
                    if loai_vb == "ND30": TienIchDocx.ve_duong_ke_shape(p, 170)
                else:
                    TienIchDocx.sao_chep_run(p, None, line, force_size=14, force_bold=False)

            cell_br = new_table.cell(1, 1)
            if loai_vb == "ND30":
                TienIchDocx.add_blank_line_if_needed(cell_br, size=14)

            for line in ngay_lines:
                p = cell_br.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.widow_control = False
                p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(0), Pt(0)
                TienIchDocx.sao_chep_run(p, None, line, force_size=13 if loai_vb=="ND30" else 14, force_bold=False, force_italic=True)

        if not is_cong_van and body_start_idx > 0:
            TienIchDocx.add_blank_line_if_needed(doc_moi, title_size)
            raw_title_text = []
            for i in range(body_start_idx):
                item_type, p_old = pre_footer_elements[i]
                if item_type == 'p':
                    lines = re.split(r'[\n\v]', p_old.text)
                    for l in lines:
                        cl = l.strip()
                        if cl and not re.match(r'^[\s\*\-\_\=\.]+$', cl): raw_title_text.append(cl)
            
            if raw_title_text:
                p_title = doc_moi.add_paragraph()
                p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                TienIchDocx.sao_chep_run(p_title, None, raw_title_text[0].upper(), force_size=title_size, force_bold=True)
                TienIchDocx.highlight_paragraph(p_title, danh_sach_tu_sai_tong)
                
                for ty_line in raw_title_text[1:]:
                    if ty_line.upper().startswith("V/V"): ty_line = "Về việc " + ty_line[3:].strip()
                    p_ty = doc_moi.add_paragraph()
                    p_ty.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    TienIchDocx.sao_chep_run(p_ty, None, ty_line, force_size=13 if loai_vb=="ND30" and shrink_mode else 14, force_bold=True)
                    TienIchDocx.highlight_paragraph(p_ty, danh_sach_tu_sai_tong)
                
                if loai_vb == "HD36":
                    p_dash = doc_moi.add_paragraph()
                    p_dash.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    TienIchDocx.sao_chep_run(p_dash, None, "-----", force_size=14, force_bold=False)
                else:
                    TienIchDocx.ve_duong_ke_shape(doc_moi.paragraphs[-1], 120)
                
                TienIchDocx.add_blank_line_if_needed(doc_moi, kg_size)

        if kinh_gui_buffer:
            p_b4 = doc_moi.add_paragraph()
            p_b4.paragraph_format.space_before, p_b4.paragraph_format.space_after = Pt(0), Pt(0)
            p_b4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            TienIchDocx.sao_chep_run(p_b4, None, " ", force_size=kg_size)

            kg_texts = []
            for p in kinh_gui_buffer:
                txt = BoXuLyVanBan.get_text_convert_bullet(p)
                if txt: kg_texts.append(txt)
                
            kg_text_full = "\n".join(kg_texts)
            is_single_kg = True if ("-" not in kg_text_full and "+" not in kg_text_full) else False

            for text in kg_texts:
                text = re.sub(r' {2,}', ' ', text).strip()
                if not text: continue
                new_p = doc_moi.add_paragraph()
                new_p.paragraph_format.widow_control = False
                new_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                new_p.paragraph_format.space_before, new_p.paragraph_format.space_after = Pt(0), Pt(0)
                
                if is_single_kg:
                    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    new_p.paragraph_format.left_indent = Cm(0)
                    new_p.paragraph_format.first_line_indent = Cm(0)
                    TienIchDocx.sao_chep_run(new_p, None, text, force_size=kg_size, force_italic=(loai_vb=="HD36"))
                else:
                    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    if text.upper().startswith("KÍNH GỬI"):
                        new_p.paragraph_format.left_indent = Cm(1.27) 
                        new_p.paragraph_format.first_line_indent = Cm(0)
                        TienIchDocx.sao_chep_run(new_p, None, text, force_size=kg_size, force_italic=(loai_vb=="HD36"))
                    elif text.startswith("-") or text.startswith("+"):
                        new_p.paragraph_format.left_indent = Cm(3.81) 
                        new_p.paragraph_format.first_line_indent = Cm(-1.27) 
                        text = re.sub(r'^[-+]\s*', '-\t', text)
                        pPr = new_p._element.get_or_add_pPr()
                        tabs = pPr.get_or_add_tabs()
                        tab = OxmlElement('w:tab')
                        tab.set(qn('w:val'), 'left'); tab.set(qn('w:pos'), str(int(3.81 * 567))) 
                        tabs.append(tab)
                        TienIchDocx.sao_chep_run(new_p, None, text, force_size=kg_size, force_italic=(loai_vb=="HD36"))
                    else:
                        new_p.paragraph_format.left_indent = Cm(3.81)
                        new_p.paragraph_format.first_line_indent = Cm(0)
                        TienIchDocx.sao_chep_run(new_p, None, text, force_size=kg_size, force_italic=(loai_vb=="HD36"))
                TienIchDocx.highlight_paragraph(new_p, danh_sach_tu_sai_tong)
            
            p_aft = doc_moi.add_paragraph()
            p_aft.paragraph_format.space_before, p_aft.paragraph_format.space_after = Pt(0), Pt(0)
            p_aft.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            TienIchDocx.sao_chep_run(p_aft, None, " ", force_size=kg_size)

        total_body_elements = len(pre_footer_elements) - body_start_idx
        for idx_offset, i in enumerate(range(body_start_idx, len(pre_footer_elements))):
            if total_body_elements > 0:
                current_p = 75.0 + (idx_offset / total_body_elements) * 15.0
                progress_callback(current_p)

            item_type, item = pre_footer_elements[i]
            if item_type == 'p':
                text_clean = BoXuLyVanBan.get_text_convert_bullet(item)
                if not text_clean: continue
                if re.match(r'^[\s\*\-\_\=\.]+$', text_clean): continue 
                
                new_p = doc_moi.add_paragraph()
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                new_p.paragraph_format.widow_control = False
                new_p.paragraph_format.left_indent = Cm(0)
                new_p.paragraph_format.first_line_indent = Mm(12.7) if loai_vb == "ND30" else Cm(1.0)
                
                if loai_vb == "HD36":
                    new_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    new_p.paragraph_format.line_spacing = Pt(18)
                else:
                    new_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                
                new_p.paragraph_format.space_before, new_p.paragraph_format.space_after = Pt(0), Pt(6)
                
                merged_runs = []
                if item.runs:
                    for r in item.runs:
                        if not r.text: continue
                        if not merged_runs:
                            merged_runs.append([r, r.text])
                        else:
                            prev_r = merged_runs[-1][0]
                            if (r.font.bold == prev_r.font.bold) and (r.font.italic == prev_r.font.italic) and (r.font.underline == prev_r.font.underline):
                                merged_runs[-1][1] += r.text
                            else:
                                merged_runs.append([r, r.text])
                    
                    is_first_text = True
                    for run_old, text_run in merged_runs:
                        text_run = text_run.replace('\n', ' ').replace('\x0b', ' ').replace('\v', ' ')
                        if is_first_text:
                            if text_run.strip(): 
                                text_run = re.sub(r'^[\s\xA0]+', '', text_run) 
                                is_first_text = False
                            else: text_run = "" 
                        
                        TienIchDocx.sao_chep_run(new_p, run_old, text_run, force_size=body_size)
                    TienIchDocx.highlight_paragraph(new_p, danh_sach_tu_sai_tong)
                else:
                    TienIchDocx.sao_chep_run(new_p, None, text_clean, force_size=body_size)
                    TienIchDocx.highlight_paragraph(new_p, danh_sach_tu_sai_tong)
                last_p_content = new_p
                
            elif item_type == 'tbl':
                new_table_xml = copy.deepcopy(item._element)
                tblPr_list = new_table_xml.xpath('./w:tblPr')
                if tblPr_list:
                    for tag in ['w:tblpPr', 'w:tblOverlap']:
                        for el in tblPr_list[0].xpath(f'./{tag}'): TienIchDocx.xoa_element(el)
                for bad_tag in new_table_xml.xpath('.//w:bookmarkStart | .//w:bookmarkEnd'): TienIchDocx.xoa_element(bad_tag)
                for node in new_table_xml.xpath('.//*[@w:id]'): node.attrib.pop(qn('w:id'), None)
                
                dummy_p = doc_moi.add_paragraph()
                dummy_p._element.addprevious(new_table_xml)
                TienIchDocx.xoa_element(dummy_p._element)
                
                if doc_moi.tables:
                    new_tbl = doc_moi.tables[-1]
                    for row in new_tbl.rows:
                        for cell in row.cells:
                            for cp in cell.paragraphs:
                                TienIchDocx.highlight_paragraph(cp, danh_sach_tu_sai_tong)
                last_p_content = None 

        progress_callback(90.0)

        while len(doc_moi.paragraphs) > 0:
            last_p = doc_moi.paragraphs[-1]
            if not last_p.text.strip() and not last_p.runs:
                TienIchDocx.xoa_element(last_p._element)
            else: break

        if last_p_content is not None and last_p_content.runs:
            last_p_content.paragraph_format.space_after = Pt(12) 
            for r in last_p_content.runs: r.font.size = Pt(body_size)
            
            clean_text = "".join([r.text for r in last_p_content.runs if r.text])
            clean_text = clean_text.rstrip(' \t\n\r./') 
            
            for r in last_p_content.runs: r.clear()
            r_clean = last_p_content.add_run(clean_text)
            TienIchDocx.thiet_lap_font(r_clean, body_size, False, False)
            
            end_char = "." if loai_vb == "HD36" else "./."
            r_end = last_p_content.add_run(end_char)
            TienIchDocx.thiet_lap_font(r_end, body_size, False, False)
        else:
            end_char = "." if loai_vb == "HD36" else "./."
            p_end = doc_moi.add_paragraph(end_char)
            p_end.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_end.paragraph_format.space_after = Pt(12)
            r_end = p_end.runs[0] if p_end.runs else p_end.add_run()
            TienIchDocx.thiet_lap_font(r_end, body_size, False, False)

        if footer_table:
            valid_rows = [row for row in footer_table.rows if "".join(cell.text for cell in row.cells).strip()]
            new_table = doc_moi.add_table(rows=len(valid_rows), cols=len(footer_table.columns))
            new_table.autofit = True
            tblPr = new_table._tbl.tblPr
            for w in tblPr.xpath('./w:tblW'): tblPr.remove(w)
            for layout in tblPr.xpath('./w:tblLayout'): tblPr.remove(layout)
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '5000'); tblW.set(qn('w:type'), 'pct'); tblPr.append(tblW)

            for row in new_table.rows:
                for c in row.cells: c.text = ""
                row.cells[0].width = Mm(90); row.cells[-1].width = Mm(70)

            for r_idx, row in enumerate(valid_rows):
                for c_idx, cell in enumerate(row.cells):
                    new_cell = new_table.cell(r_idx, c_idx)
                    lines = [l.strip() for l in cell.text.split('\n') if l.strip()]
                    if not lines: new_cell.add_paragraph().paragraph_format.widow_control = False; continue

                    p_first_used = False
                    for i, line in enumerate(lines):
                        if not p_first_used:
                            p = new_cell.paragraphs[0]; p_first_used = True
                        else: p = new_cell.add_paragraph()
                            
                        p.paragraph_format.widow_control = False
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(0), Pt(0)
                        line_up = line.upper()

                        if c_idx == 0: 
                            if "NƠI NHẬN" in line_up:
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
                                if loai_vb == "HD36": TienIchDocx.copy_text_only_with_subscript(p, line, 14, bold=False, italic=False, underline=True)
                                else: TienIchDocx.copy_text_only_with_subscript(p, line, 12, bold=True, italic=True, underline=False)
                            else:
                                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
                                text_clean = re.sub(r'^[-+]\s*', '- ', line)
                                if loai_vb == "HD36": TienIchDocx.copy_text_only_with_subscript(p, text_clean, 12, bold=False, italic=False)
                                else: TienIchDocx.copy_text_only_with_subscript(p, text_clean, 11, bold=False, italic=False)
                        elif c_idx > 0: 
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            is_name = (i == len(lines) - 1)
                            if loai_vb == "HD36":
                                if is_name: TienIchDocx.copy_text_only_with_subscript(p, line, name_size, bold=True)
                                else:
                                    if any(kw in line_up for kw in ["T/M", "TM.", "TL.", "KT.", "Q.", "THAY MẶT"]): TienIchDocx.copy_text_only_with_subscript(p, line_up, sig_size, bold=True)
                                    else: TienIchDocx.copy_text_only_with_subscript(p, line_up, sig_size, bold=False)
                            else:
                                TienIchDocx.copy_text_only_with_subscript(p, line, sig_size, bold=True)
                            
                            if any(kw in line_up for kw in ["CHỦ TỊCH", "TM.", "KT.", "BÍ THƯ", "T/M"]):
                                is_last_title = True
                                if i + 1 < len(lines) and any(kw in lines[i+1].upper() for kw in ["CHỦ TỊCH", "TM.", "KT.", "BÍ THƯ", "T/M"]):
                                    is_last_title = False
                                if is_last_title: TienIchDocx.tao_dong_trang_chu_ky(new_cell, 6, size=sig_size)

        if danh_sach_tu_sai_tong:
            ds_unique = list(set([t for t in danh_sach_tu_sai_tong]))
            log_callback(f"<b><span style='color:#d35400;'>➜ HỆ THỐNG ĐÃ TỰ ĐỘNG BÔI ĐỎ CÁC CỤM TỪ SAU: {', '.join(ds_unique)}</span></b>", "#d35400")

        doc_moi.save(output_path)
        progress_callback(92.0)

    @staticmethod
    def check_orphan_signature(docx_path):
        need_resize = False
        try:
            pythoncom.CoInitialize()
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            abs_path = os.path.abspath(docx_path)
            doc = word.Documents.Open(abs_path)
            
            if doc.ComputeStatistics(2) <= 2:
                for sec in doc.Sections:
                    for header in sec.Headers:
                        if header.Exists: header.Range.Text = ""
            
            if doc.Tables.Count > 0:
                for i in range(doc.Tables.Count, 0, -1):
                    tbl = doc.Tables(i)
                    txt = tbl.Range.Text.upper()
                    if "NƠI NHẬN" in txt or "CHỦ TỊCH" in txt or "TM." in txt:
                        start_page = tbl.Range.Information(3) 
                        last_text_pos = 0
                        for p in doc.Paragraphs:
                            if p.Range.Start >= tbl.Range.Start: break
                            if len(p.Range.Text.strip()) > 0:
                                last_text_pos = p.Range.Start
                        if last_text_pos > 0:
                            last_text_page = doc.Range(last_text_pos, last_text_pos).Information(3)
                            if start_page > last_text_page:
                                need_resize = True
                        break
            doc.Close(SaveChanges=True)
            word.Quit()
        except Exception: pass
        finally: pythoncom.CoUninitialize()
        return need_resize

    @staticmethod
    def process_with_two_pass(input_path, log_callback, progress_callback, enable_spellcheck=True):
        output_temp = os.path.join(os.path.dirname(input_path), "TEMP_" + os.path.basename(input_path))
        if output_temp.endswith('.doc'): output_temp += "x"
        
        doc_type = BoXuLyVanBan.phan_loai_van_ban(Document(input_path))
        
        BoXuLyVanBan.process_single_pass(input_path, output_temp, log_callback, progress_callback, enable_spellcheck=enable_spellcheck, shrink_mode=False)
        progress_callback(95.0)
        
        output_final = os.path.join(os.path.dirname(input_path), ("HD36_" if doc_type == "HD36" else "ND30_") + os.path.basename(input_path))
        if output_final.endswith('.doc'): output_final += "x"

        if doc_type == "ND30":
            log_callback("<i>➔ Kiểm tra độ bám dính của Chữ ký & Nơi nhận cuối trang...</i>", "#7f8c8d")
            need_resize = BoXuLyVanBan.check_orphan_signature(output_temp)
            if need_resize:
                log_callback("<i>➜ Phát hiện rớt Chữ ký sang trang trắng: Đang tự động Vượt 2 co văn bản...</i>", "#e67e22")
                BoXuLyVanBan.process_single_pass(input_path, output_final, lambda t,c: None, progress_callback, enable_spellcheck=False, shrink_mode=True)
                if os.path.exists(output_temp): os.remove(output_temp)
                BoXuLyVanBan.check_orphan_signature(output_final) 
                log_callback("<b>➜ Đã fix chống rớt trang thành công!</b>", "#27ae60")
                progress_callback(99.0)
                return output_final
            else:
                log_callback("<i>➜ Bố cục an toàn, không bị rớt trang.</i>", "#27ae60")
        
        try:
            pythoncom.CoInitialize()
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            doc = word.Documents.Open(os.path.abspath(output_temp))
            if doc.ComputeStatistics(2) <= 2:
                for sec in doc.Sections:
                    for header in sec.Headers:
                        if header.Exists: header.Range.Text = ""
            doc.Close(SaveChanges=True)
            word.Quit()
        except: pass
        finally: pythoncom.CoUninitialize()
        
        if os.path.exists(output_final): os.remove(output_final)
        os.rename(output_temp, output_final)
        progress_callback(100.0)
        return output_final

class DocumentWorker(QThread):
    progress_updated = pyqtSignal(float)
    log_updated = pyqtSignal(str, str)
    file_finished = pyqtSignal(str, str) 
    all_finished = pyqtSignal()
    preview_ready = pyqtSignal(str) 

    def __init__(self, files, enable_spellcheck):
        super().__init__()
        self.files = files
        self.enable_spellcheck = enable_spellcheck
        self.global_file_counter = 0

    def run(self):
        pythoncom.CoInitialize()
        for idx, f in enumerate(self.files):
            self.global_file_counter += 1
            if f.endswith('.docx') or f.endswith('.doc'):
                try:
                    self.progress_updated.emit(5.0)
                    file_xu_ly = f
                    is_converted = False
                    
                    self.log_updated.emit(f"<b>--- [ĐANG XỬ LÝ {self.global_file_counter}/{len(self.files)}] : {os.path.basename(f)} ---</b>", "#2c3e50")
                    
                    if f.endswith('.doc'):
                        try:
                            file_xu_ly = TienIchDocx.convert_doc_to_docx(f)
                            is_converted = True
                        except Exception as e:
                            self.log_updated.emit(f"<b>Lỗi Convert:</b> {str(e)}", "#e74c3c")
                            continue

                    output_path = BoXuLyVanBan.process_with_two_pass(
                        file_xu_ly, 
                        lambda text, color: self.log_updated.emit(text, color), 
                        lambda val: self.progress_updated.emit(val),
                        self.enable_spellcheck
                    )
                    
                    if is_converted and os.path.exists(file_xu_ly): os.remove(file_xu_ly)
                    self.log_updated.emit("<b>➜ CHUẨN HÓA THÀNH CÔNG! Đang xuất bản xem trước (Preview)...</b>", "#27ae60")
                    
                    pdf_preview_path = BoXuLyVanBan.export_preview_pdf(output_path)
                    self.preview_ready.emit(pdf_preview_path if pdf_preview_path else "")

                    self.file_finished.emit(output_path, f)
                except Exception as e:
                    self.log_updated.emit(f"<b>LỖI HỆ THỐNG:</b> {str(e)}", "#c0392b")
            else:
                self.log_updated.emit(f"Cảnh báo: File {os.path.basename(f)} không phải Word", "#d35400")
        
        pythoncom.CoUninitialize()
        self.all_finished.emit()

class ActionButtonPanel(QWidget):
    def __init__(self, new_file, old_file, parent=None):
        super().__init__(parent)
        self.new_file = new_file
        self.old_file = old_file
        
        layout = QHBoxLayout(self)
        layout.setContentsMargins(10, 5, 10, 5)
        layout.setSpacing(10)
        
        btn_new = QPushButton("📄 MỞ FILE MỚI")
        btn_new.setStyleSheet(self.get_btn_style("#3498db", "#2980b9"))
        btn_new.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_new.clicked.connect(self.mo_file_moi)
        
        btn_old = QPushButton("📄 MỞ FILE GỐC")
        btn_old.setStyleSheet(self.get_btn_style("#95a5a6", "#7f8c8d"))
        btn_old.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_old.clicked.connect(self.mo_file_goc)
        
        btn_folder = QPushButton("📂 MỞ THƯ MỤC")
        btn_folder.setStyleSheet(self.get_btn_style("#e67e22", "#d35400"))
        btn_folder.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_folder.clicked.connect(self.mo_thu_muc)

        layout.addWidget(btn_new)
        layout.addWidget(btn_old)
        layout.addWidget(btn_folder)

    def get_btn_style(self, top_color, bottom_color):
        return f"""
        QPushButton {{
            background-color: {top_color}; color: white; border-radius: 6px;
            border-top: 1px solid #ffffff; border-left: 1px solid #ffffff;
            border-right: 2px solid {bottom_color}; border-bottom: 4px solid {bottom_color};
            font-weight: bold; font-size: 13px; padding: 8px;
        }}
        QPushButton:hover {{ background-color: {bottom_color}; }}
        QPushButton:pressed {{
            background-color: {bottom_color}; border-top: 4px solid transparent; border-bottom: 0px solid transparent;
            border-left: 2px solid transparent; border-right: 0px solid transparent; margin-top: 2px;
        }}
        """

    def open_local_file(self, file_path):
        if sys.platform == "win32":
            try: os.startfile(os.path.normpath(file_path))
            except Exception: QDesktopServices.openUrl(QUrl.fromLocalFile(file_path))
        elif sys.platform == "darwin": subprocess.call(["open", file_path])
        else: subprocess.call(["xdg-open", file_path])

    def mo_file_moi(self): self.open_local_file(self.new_file)
    def mo_file_goc(self): self.open_local_file(self.old_file)
    def mo_thu_muc(self):
        if os.path.exists(self.new_file):
            if sys.platform == "win32": subprocess.Popen(f'explorer /select,"{os.path.normpath(self.new_file)}"')
            elif sys.platform == "darwin": subprocess.Popen(["open", "-R", self.new_file])
            else: subprocess.Popen(["xdg-open", os.path.dirname(self.new_file)])

class BangLamViec(QMainWindow):
    def __init__(self):
        super().__init__()
        version = datetime.datetime.now().strftime("Master Hybrid V30.%d.%m")
        self.setWindowTitle(f"Trợ lý Văn Thư Kép AI - UBND xã Đồng Tiến ({version})")
        # Phục hồi các nút điều khiển cửa sổ tiêu chuẩn của Windows (Mở rộng, Thu nhỏ, Tắt)
        self.setWindowFlags(Qt.WindowType.Window | Qt.WindowType.WindowMinimizeButtonHint | Qt.WindowType.WindowMaximizeButtonHint | Qt.WindowType.WindowCloseButtonHint)
        self.resize(1300, 800)
        self.setMinimumSize(1000, 600)
        self.setAcceptDrops(True)
        
        self.total_files = 0
        self.global_success_counter = 0
        self.start_time = 0
        self.worker = None

        main_splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # --- CỘT TRÁI (Công cụ & Log) ---
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(10, 10, 10, 10)
        
        self.btn_browse = QPushButton("🔍 CHỌN FILE TỪ MÁY TÍNH (.doc/.docx)")
        self.btn_browse.setStyleSheet("""
        QPushButton {
            background-color: #2ecc71; color: white; border-radius: 8px;
            border-bottom: 4px solid #27ae60; font-weight: bold; font-size: 16px; padding: 15px;
        }
        QPushButton:hover { background-color: #27ae60; }
        QPushButton:disabled { background-color: #95a5a6; border: none; }
        """)
        self.btn_browse.clicked.connect(self.browse_files)
        
        pdf_layout = QHBoxLayout()
        self.btn_nd30 = QPushButton("📖 Xem NĐ 30")
        self.btn_nd30.setStyleSheet(self.get_pdf_btn_style("#9b59b6", "#8e44ad"))
        self.btn_nd30.clicked.connect(self.open_nd30_pdf)
        self.btn_hd36 = QPushButton("📖 Xem HD 36")
        self.btn_hd36.setStyleSheet(self.get_pdf_btn_style("#9b59b6", "#8e44ad"))
        self.btn_hd36.clicked.connect(self.open_hd36_pdf)
        pdf_layout.addWidget(self.btn_nd30)
        pdf_layout.addWidget(self.btn_hd36)

        self.chk_spellcheck = QCheckBox("Bật rà soát chính tả (AI + Từ điển)")
        self.chk_spellcheck.setChecked(True)
        self.chk_spellcheck.setStyleSheet("font-size: 14px; font-weight: bold; color: #2c3e50; margin-top: 5px;")
        
        self.lbl_stats = QLabel("<b>Thống kê: Chưa có file nào được xử lý</b>")
        self.lbl_stats.setStyleSheet("font-size: 14px; color: #2980b9; padding: 5px; background: #e8f4f8; border-radius: 5px;")
        self.lbl_stats.setAlignment(Qt.AlignmentFlag.AlignCenter)

        self.progress_bar = QProgressBar()
        self.progress_bar.setStyleSheet("QProgressBar { border: 2px solid #bdc3c7; border-radius: 5px; text-align: center; font-weight: bold; } QProgressBar::chunk { background-color: #3498db; }")
        self.progress_bar.setValue(0)
        self.progress_bar.hide()

        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_content = QWidget()
        self.scroll_layout = QVBoxLayout(self.scroll_content)
        self.scroll_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.title_log = QLabel("<h3 style='color:#2c3e50; text-align:center;'>LOG CHI TIẾT</h3>")
        self.scroll_layout.addWidget(self.title_log)
        self.scroll_area.setWidget(self.scroll_content)

        credit_text = "Mọi quan tâm, báo lỗi xin gửi đến: nguyenduccong@me.com"
        self.lbl_credit = QLabel(credit_text)
        self.lbl_credit.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lbl_credit.setStyleSheet("color: #7f8c8d; font-size: 12px; font-weight: bold; font-style: italic; margin-top: 5px;")

        left_layout.addWidget(self.btn_browse)
        left_layout.addLayout(pdf_layout)
        left_layout.addWidget(self.chk_spellcheck)
        left_layout.addWidget(self.lbl_stats)
        left_layout.addWidget(self.progress_bar)
        left_layout.addWidget(self.scroll_area, stretch=1)
        left_layout.addWidget(self.lbl_credit)
        
        # --- CỘT PHẢI (Preview Bằng PDFViewer 100% Zoomable) ---
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(0, 10, 10, 10)
        
        lbl_preview_title = QLabel("<h3>XEM TRƯỚC VĂN BẢN (100% CHUẨN MS WORD)</h3>")
        lbl_preview_title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl_preview_title.setStyleSheet("color: #d35400; background: #fae5d3; border-radius: 5px; padding: 5px;")
        
        self.pdf_viewer = PDFViewer()
        
        right_layout.addWidget(lbl_preview_title)
        right_layout.addWidget(self.pdf_viewer)

        main_splitter.addWidget(left_widget)
        main_splitter.addWidget(right_widget)
        main_splitter.setSizes([400, 900])

        self.setCentralWidget(main_splitter)

    def get_pdf_btn_style(self, top_color, bottom_color):
        return f"QPushButton {{ background-color: {top_color}; color: white; border-radius: 6px; font-weight: bold; font-size: 13px; padding: 8px 10px; border-bottom: 4px solid {bottom_color}; }} QPushButton:hover {{ background-color: {bottom_color}; }}"

    def open_pdf_file(self, filename):
        file_path = resource_path(filename)
        if os.path.exists(file_path):
            if sys.platform == "win32": os.startfile(os.path.normpath(file_path))
            elif sys.platform == "darwin": subprocess.call(["open", file_path])
            else: subprocess.call(["xdg-open", file_path])
        else: QMessageBox.warning(self, "Lỗi", f"Không tìm thấy tài liệu:\n{filename}")

    def open_nd30_pdf(self): self.open_pdf_file("30.signed.pdf")
    def open_hd36_pdf(self): self.open_pdf_file("hd36vptw-01 (3).pdf")

    def append_log(self, text, color="#34495e"):
        lbl = QLabel(f"<span style='color:{color}; font-size: 13px;'>{text}</span>")
        lbl.setWordWrap(True)
        self.scroll_layout.addWidget(lbl)
        self.scroll_area.verticalScrollBar().setValue(self.scroll_area.verticalScrollBar().maximum())

    def update_progress(self, val_float):
        self.progress_bar.setRange(0, 10000)
        self.progress_bar.setValue(int(val_float * 100))
        self.progress_bar.setFormat(f"{val_float:.2f} %")
        
    def add_button_panel(self, output_path, input_path):
        btn_panel = ActionButtonPanel(output_path, input_path, self)
        self.scroll_layout.addWidget(btn_panel)
        self.global_success_counter += 1

    def finish_processing(self):
        self.update_progress(100.0)
        self.progress_bar.hide()
        self.btn_browse.setEnabled(True)
        elapsed = round(time.time() - self.start_time, 1)
        self.lbl_stats.setText(f"<b>Thống kê: Đã hoàn tất {self.global_success_counter}/{self.total_files} file | Thời gian: {elapsed} giây</b>")

    def browse_files(self):
        file_paths, _ = QFileDialog.getOpenFileNames(self, "Chọn Văn Bản", "", "Word (*.doc *.docx)")
        if file_paths: self.process_files(file_paths)

    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls(): event.accept()

    def dropEvent(self, event: QDropEvent):
        files = [u.toLocalFile() for u in event.mimeData().urls()]
        self.process_files(files)

    def process_files(self, files):
        if not files: return
        self.btn_browse.setEnabled(False)
        self.progress_bar.show()
        self.progress_bar.setValue(0)
        
        self.pdf_viewer.show_message("Đang rà soát và xử lý văn bản, vui lòng đợi...")
        
        self.total_files = len(files)
        self.global_success_counter = 0
        self.start_time = time.time()
        
        enable_sc = self.chk_spellcheck.isChecked()
        
        self.worker = DocumentWorker(files, enable_sc)
        self.worker.progress_updated.connect(self.update_progress)
        self.worker.log_updated.connect(self.append_log)
        self.worker.preview_ready.connect(self.pdf_viewer.load_pdf) # Truyền thẳng file vào bộ Zoom
        self.worker.file_finished.connect(self.add_button_panel)
        self.worker.all_finished.connect(self.finish_processing)
        self.worker.start()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = BangLamViec()
    window.show()
    sys.exit(app.exec())